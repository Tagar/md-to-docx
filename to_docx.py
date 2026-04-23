#!/usr/bin/env python3
"""Convert plain-text / light-markdown drafts into nicely-formatted Word .docx files.

Usage:
    to_docx.py INPUT.txt [OUTPUT.docx] [--title "Document title"]

Header block:
    - First non-empty line                                      Title (Word Title style)
    - Second non-empty plain-text line (short, Title-Case-ish)  Subtitle (Word Subtitle style)
    - Subsequent lines                                          body

Markdown features supported:
    Block-level:
    - `# H1`, `## H2`, `### H3`, `#### H4`                     explicit headings
    - `ALL CAPS LINE` or `1. ALL CAPS TITLE`                   implicit H1
    - Short `1. Title Case` (no end punctuation)               implicit H2
    - `N. Sentence with punctuation.`                          numbered list item
    - `- ` or `* ` (indent = sub-bullet)                       bullet
    - `> quoted text`                                          blockquote (multi-line ok)
    - ` ``` ` fenced code block                                monospace paragraph
    - ` ```mermaid ` fenced block                              rendered via `mmdc` (if installed), else code
    - Pipe tables (`| a | b |` + separator row)                docx table
    - `![alt](path "width=6in")` on its own line               embedded image + italic caption
    - Divider lines (`====` / `----`)                          dropped

    Inline (inside any paragraph/bullet/list/cell/heading):
    - `**text**` or `__text__`                                 bold
    - `*text*` or `_text_`                                     italic
    - `` `text` ``                                             monospace / inline code
    - `~~text~~`                                               strikethrough
    - `[label](url)`                                           hyperlink
    - `<html>` tags                                            stripped

    Label lines `Label: value`:                                bold label + value

Unit tests live in tests/test_to_docx.py.
"""
from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ----- inline tokenization -----

# Order matters: `**` before `*`, `__` before `_`, tokens compete left-to-right.
INLINE_RE = re.compile(
    r"\*\*(?P<bold_ast>.+?)\*\*"                        # **bold**
    r"|__(?P<bold_und>.+?)__"                            # __bold__
    r"|\*(?P<italic_ast>[^*\n]+?)\*"                     # *italic*  (not **)
    r"|(?<![A-Za-z0-9_])_(?P<italic_und>[^_\n]+?)_(?![A-Za-z0-9_])"  # _italic_
    r"|`(?P<code>[^`\n]+?)`"                             # `inline code`
    r"|~~(?P<strike>.+?)~~"                              # ~~strike~~
    r"|\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)\s]+)\)"  # [text](url)
)

HTML_TAG_RE = re.compile(r"<!--.*?-->|<[^>]+>", re.DOTALL)


def strip_html(text: str) -> str:
    return HTML_TAG_RE.sub("", text)


def _add_hyperlink(paragraph, url: str, text: str):
    """Attach a clickable hyperlink run to `paragraph`."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _style_code_run(run) -> None:
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)  # dusty-rose code color


def add_inline_runs(paragraph, text: str) -> None:
    """Add `text` to `paragraph`, parsing markdown inline tokens into styled runs."""
    if not text:
        return
    text = strip_html(text)
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])

        if m.group("bold_ast") is not None:
            paragraph.add_run(m.group("bold_ast")).bold = True
        elif m.group("bold_und") is not None:
            paragraph.add_run(m.group("bold_und")).bold = True
        elif m.group("italic_ast") is not None:
            paragraph.add_run(m.group("italic_ast")).italic = True
        elif m.group("italic_und") is not None:
            paragraph.add_run(m.group("italic_und")).italic = True
        elif m.group("code") is not None:
            run = paragraph.add_run(m.group("code"))
            _style_code_run(run)
        elif m.group("strike") is not None:
            run = paragraph.add_run(m.group("strike"))
            run.font.strike = True
        elif m.group("link_text") is not None:
            _add_hyperlink(paragraph, m.group("link_url"), m.group("link_text"))
        pos = m.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])


# ----- block-level classifiers -----

def is_divider(line: str) -> bool:
    s = line.strip()
    return len(s) >= 8 and set(s) <= {"=", "-"}


def is_markdown_heading(line: str):
    m = re.match(r"^(#{1,4})\s+(.+?)\s*$", line)
    if not m:
        return None
    return len(m.group(1)), m.group(2)


def is_all_caps_heading(line: str) -> bool:
    s = line.strip()
    if not s or len(s) < 3 or len(s) > 80:
        return False
    letters = [c for c in s if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters)


def is_numbered_heading(line: str) -> bool:
    stripped = line.strip()
    return (
        bool(re.match(r"^\d+\.\s+[A-Z]", stripped))
        and len(stripped) < 100
        and not stripped.endswith(("?", ".", ",", ";", ":"))
    )


def is_numbered_list_item(line: str) -> bool:
    return bool(re.match(r"^\s*\d+\.\s+", line))


def is_bullet(line: str) -> bool:
    stripped = line.lstrip()
    return stripped.startswith("- ") or stripped.startswith("* ")


def is_kv_label(line: str) -> bool:
    return bool(re.match(r"^[A-Z][a-zA-Z ]{0,25}:\s", line.strip()))


def is_blockquote(line: str) -> bool:
    return line.lstrip().startswith(">")


def is_code_fence(line: str) -> bool:
    return line.strip().startswith("```")


MERMAID_FENCE_RE = re.compile(r"^\s*```\s*mermaid\b", re.IGNORECASE)


def is_mermaid_fence(line: str) -> bool:
    """True if this is an opening fence for a Mermaid diagram (```mermaid)."""
    return bool(MERMAID_FENCE_RE.match(line))


def is_table_row(line: str) -> bool:
    return line.lstrip().startswith("|") and line.rstrip().endswith("|")


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", stripped)) and "-" in stripped and "|" in stripped


IMAGE_BLOCK_RE = re.compile(
    r'^\s*!\[(?P<alt>[^\]]*)\]\((?P<path>[^)\s"]+)(?:\s+"(?P<title>[^"]*)")?\)\s*$'
)


def is_image_block(line: str) -> bool:
    return bool(IMAGE_BLOCK_RE.match(line))


def parse_image_block(line: str):
    """Return (alt, path, title) or None if not an image block."""
    m = IMAGE_BLOCK_RE.match(line)
    if not m:
        return None
    return m.group("alt"), m.group("path"), m.group("title")


def is_subtitle_candidate(line: str) -> bool:
    """A plain-text line that should render as Word Subtitle style.

    Used only for the position immediately after the document title — the
    first non-empty, non-header line at the top of the file. Must look like
    a title/subtitle: short, starts uppercase, no end punctuation, no
    block-level markers, <=10 words.
    """
    stripped = line.strip()
    if not stripped or len(stripped) > 80:
        return False
    if stripped[-1] in ".!?;,:":
        return False
    if len(stripped.split()) > 10:
        return False
    if not stripped[0].isupper():
        return False
    if (is_markdown_heading(line) or is_divider(line) or is_bullet(line)
            or is_numbered_list_item(line) or is_kv_label(line)
            or is_blockquote(line) or is_code_fence(line) or is_table_row(line)
            or is_all_caps_heading(line) or is_numbered_heading(line)):
        return False
    return True


def parse_table_row(line: str) -> list[str]:
    # Strip leading/trailing pipes, split on inner pipes, strip each cell
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    return [c.strip() for c in inner.split("|")]


# ----- block builders -----

def _add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell = table.cell(r_idx, c_idx)
            cell.paragraphs[0].clear()
            text = row[c_idx] if c_idx < len(row) else ""
            para = cell.paragraphs[0]
            add_inline_runs(para, text)
            # Header row: bold all runs
            if r_idx == 0:
                for run in para.runs:
                    run.bold = True


def _add_blockquote(doc, lines: list[str]) -> None:
    text = " ".join(l.strip() for l in lines)
    try:
        p = doc.add_paragraph(style="Intense Quote")
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
    add_inline_runs(p, text)


def _add_code_block(doc, code: str) -> None:
    for code_line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        run = p.add_run(code_line if code_line else " ")
        _style_code_run(run)


def _add_subtitle(doc, text: str) -> None:
    try:
        p = doc.add_paragraph(style="Subtitle")
    except KeyError:
        p = doc.add_paragraph()
    add_inline_runs(p, text)


def _render_mermaid_to_png(mermaid_source: str, temp_dir: Path):
    """Render Mermaid source to PNG via the `mmdc` CLI.

    Returns the PNG path, or None if mmdc is not installed or rendering fails.
    """
    mmdc = shutil.which("mmdc")
    if not mmdc:
        return None

    in_path = temp_dir / "diagram.mmd"
    out_path = temp_dir / "diagram.png"
    in_path.write_text(mermaid_source)

    try:
        subprocess.run(
            [mmdc, "-i", str(in_path), "-o", str(out_path), "-b", "transparent"],
            check=True,
            capture_output=True,
            timeout=60,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError):
        return None

    return out_path if out_path.exists() else None


def _add_mermaid(doc, mermaid_source: str) -> None:
    """Render a ```mermaid block to an embedded image, or fall back to a code block.

    If mmdc is unavailable or fails, prepends an italic note so readers know
    the source was meant to be a diagram and inserts the source as a normal
    code block — preserving information rather than dropping it.
    """
    with tempfile.TemporaryDirectory() as td:
        png_path = _render_mermaid_to_png(mermaid_source, Path(td))
        if png_path is None:
            note = doc.add_paragraph()
            run = note.add_run(
                "[Mermaid diagram — install @mermaid-js/mermaid-cli (`mmdc`) to render inline]"
            )
            run.italic = True
            run.font.size = Pt(9)
            _add_code_block(doc, mermaid_source)
            return

        try:
            doc.add_picture(str(png_path), width=Inches(6))
        except Exception as e:
            err = doc.add_paragraph()
            err.add_run(f"[mermaid render error: {e}]").italic = True
            _add_code_block(doc, mermaid_source)


def _add_image(doc, path_str: str, base_dir: Path, alt: str, title) -> None:
    """Insert a block-level image. Path is resolved relative to base_dir if not absolute.

    Title syntax "width=Xin" overrides the default 6-inch width.
    Missing files render as an italic placeholder paragraph instead of crashing.
    """
    path = Path(path_str)
    if not path.is_absolute():
        path = base_dir / path

    width = Inches(6)
    if title:
        m = re.search(r"width\s*=\s*([\d.]+)\s*in", title)
        if m:
            width = Inches(float(m.group(1)))

    if not path.exists():
        p = doc.add_paragraph()
        p.add_run(f"[image not found: {path_str}]").italic = True
        return

    try:
        doc.add_picture(str(path), width=width)
    except Exception as e:
        p = doc.add_paragraph()
        p.add_run(f"[image error: {e}]").italic = True
        return

    if alt:
        caption = doc.add_paragraph()
        run = caption.add_run(alt)
        run.italic = True
        run.font.size = Pt(9)


# ----- main conversion loop -----

def convert(src: Path, dst: Path, title: str) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_heading(title, level=0)

    lines = src.read_text().splitlines()
    i = 0

    # Header block: skip the first non-empty line if it matches the title
    # (default_title uses line 1, so otherwise it'd render twice).
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines) and lines[i].strip() == title:
        i += 1

    # If the next non-empty line is a subtitle candidate (short, plain text,
    # no end punctuation, Title-Case-ish), render it with the Subtitle style.
    j = i
    while j < len(lines) and not lines[j].strip():
        j += 1
    if j < len(lines) and is_subtitle_candidate(lines[j]):
        _add_subtitle(doc, lines[j].strip())
        i = j + 1

    while i < len(lines):
        line = lines[i]

        # Divider or blank -> skip
        if is_divider(line) or not line.strip():
            i += 1
            continue

        # Fenced code block (including ```mermaid → rendered diagram)
        if is_code_fence(line):
            is_mermaid = is_mermaid_fence(line)
            i += 1
            buf: list[str] = []
            while i < len(lines) and not is_code_fence(lines[i]):
                buf.append(lines[i])
                i += 1
            body = "\n".join(buf)
            if is_mermaid:
                _add_mermaid(doc, body)
            else:
                _add_code_block(doc, body)
            if i < len(lines):
                i += 1  # skip closing fence
            continue

        # Block-level image: ![alt](path "width=Xin")
        if is_image_block(line):
            alt, path, img_title = parse_image_block(line)
            _add_image(doc, path, src.parent, alt, img_title)
            i += 1
            continue

        # Pipe table: first row + separator row pattern
        if is_table_row(line) and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [parse_table_row(line)]
            i += 2  # skip header + separator
            while i < len(lines) and is_table_row(lines[i]):
                rows.append(parse_table_row(lines[i]))
                i += 1
            _add_table(doc, rows)
            continue

        # Blockquote (can span multiple > lines)
        if is_blockquote(line):
            quote_lines: list[str] = []
            while i < len(lines) and is_blockquote(lines[i]):
                quote_lines.append(lines[i].lstrip().lstrip(">").lstrip())
                i += 1
            _add_blockquote(doc, quote_lines)
            continue

        # Explicit markdown heading
        md = is_markdown_heading(line)
        if md:
            level, heading_text = md
            h = doc.add_heading("", level=level)
            add_inline_runs(h, heading_text)
            i += 1
            continue

        # Implicit ALL CAPS / all-caps-numbered top heading
        if is_all_caps_heading(line):
            doc.add_heading(line.strip().title(), level=1)
            i += 1
            continue

        # Short Title-Case numbered heading
        if is_numbered_heading(line):
            h = doc.add_heading("", level=2)
            add_inline_runs(h, line.strip())
            i += 1
            continue

        # Numbered list item
        if is_numbered_list_item(line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
            add_inline_runs(p, text)
            i += 1
            continue

        # Bullet
        if is_bullet(line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            style = "List Bullet 2" if line.startswith(("  - ", "    - ")) else "List Bullet"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph()
                p.add_run("• ")
            add_inline_runs(p, text)
            i += 1
            continue

        # `Label: value`
        if is_kv_label(line):
            label, _, rest = line.strip().partition(":")
            p = doc.add_paragraph()
            p.add_run(label + ":").bold = True
            if rest.strip():
                add_inline_runs(p, " " + rest.strip())
            i += 1
            continue

        # Default paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, line)
        i += 1

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)
    print(f"wrote {dst} ({dst.stat().st_size} bytes)")


def default_title(src: Path) -> str:
    for line in src.read_text().splitlines():
        if line.strip():
            return line.strip()
    return src.stem


def main() -> None:
    ap = argparse.ArgumentParser(description="Convert plain-text drafts to .docx")
    ap.add_argument("input", type=Path, help="input .txt or .md file")
    ap.add_argument("output", type=Path, nargs="?", default=None,
                    help="output .docx (default: <input_dir>/<stem>.docx)")
    ap.add_argument("--title", default=None,
                    help="document title (default: first non-empty line of input)")
    args = ap.parse_args()

    if not args.input.exists():
        sys.exit(f"error: {args.input} does not exist")

    out = args.output or args.input.with_suffix(".docx")
    title = args.title or default_title(args.input)
    convert(args.input, out, title)


if __name__ == "__main__":
    main()

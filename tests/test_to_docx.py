"""Unit + integration tests for scripts/to_docx.py.

Run with:
    uv run --with python-docx --with pytest pytest scripts/tests/test_to_docx.py -v
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
import to_docx  # noqa: E402


# ============================================================
# Classifier unit tests
# ============================================================

class TestClassifiers:
    @pytest.mark.parametrize("line,expected", [
        ("========", True),
        ("--------", True),
        ("========================================================", True),
        ("==-==-==-", True),
        ("===", False),       # too short
        ("", False),
        ("hello", False),
        ("== hello ==", False),
    ])
    def test_is_divider(self, line, expected):
        assert to_docx.is_divider(line) is expected

    @pytest.mark.parametrize("line,expected", [
        ("# Heading", (1, "Heading")),
        ("## Two", (2, "Two")),
        ("### Three", (3, "Three")),
        ("#### Four", (4, "Four")),
        ("##### Five", None),    # > 4 hashes not supported
        ("Not a heading", None),
        ("#notaheading", None),  # missing space
    ])
    def test_is_markdown_heading(self, line, expected):
        assert to_docx.is_markdown_heading(line) == expected

    @pytest.mark.parametrize("line,expected", [
        ("HELLO WORLD", True),
        ("1. OVERVIEW", True),
        ("2. PROJECT CONTEXT", True),
        ("Hello World", False),
        ("", False),
        ("AB", False),  # too short (<3)
    ])
    def test_is_all_caps_heading(self, line, expected):
        assert to_docx.is_all_caps_heading(line) is expected

    @pytest.mark.parametrize("line,expected", [
        ("1. Overview", True),
        ("2. Project Context", True),
        ("1. sentence ending with period.", False),
        ("1. question?", False),
        ("1. Colon:", False),
        ("Not numbered", False),
    ])
    def test_is_numbered_heading(self, line, expected):
        assert to_docx.is_numbered_heading(line) is expected

    @pytest.mark.parametrize("line,expected", [
        ("1. First item", True),
        ("  2. Indented item", True),
        ("- bullet", False),
        ("abc", False),
    ])
    def test_is_numbered_list_item(self, line, expected):
        assert to_docx.is_numbered_list_item(line) is expected

    @pytest.mark.parametrize("line,expected", [
        ("- bullet", True),
        ("* bullet", True),
        ("  - sub-bullet", True),
        ("no bullet", False),
        ("-nospace", False),
    ])
    def test_is_bullet(self, line, expected):
        assert to_docx.is_bullet(line) is expected

    @pytest.mark.parametrize("line,expected", [
        ("Label: value", True),
        ("Description: something", True),
        ("lowercase: no", False),
        ("No colon", False),
    ])
    def test_is_kv_label(self, line, expected):
        assert to_docx.is_kv_label(line) is expected

    @pytest.mark.parametrize("line,expected", [
        ("> quote", True),
        ("  > indented quote", True),
        ("regular text", False),
        (">>nested", True),
    ])
    def test_is_blockquote(self, line, expected):
        assert to_docx.is_blockquote(line) is expected

    @pytest.mark.parametrize("line,expected", [
        ("```", True),
        ("```python", True),
        ("`inline`", False),
        ("regular", False),
    ])
    def test_is_code_fence(self, line, expected):
        assert to_docx.is_code_fence(line) is expected

    @pytest.mark.parametrize("line,expected", [
        ("| col1 | col2 |", True),
        ("|--|--|", True),
        ("| single |", True),
        ("no pipe", False),
        ("| missing trailing", False),
    ])
    def test_is_table_row(self, line, expected):
        assert to_docx.is_table_row(line) is expected

    @pytest.mark.parametrize("line,expected", [
        ("|---|---|", True),
        ("| --- | --- |", True),
        ("|:---:|:---|", True),
        ("| col1 | col2 |", False),  # no dashes
        ("----", False),              # no pipes
    ])
    def test_is_table_separator(self, line, expected):
        assert to_docx.is_table_separator(line) is expected

    def test_parse_table_row(self):
        assert to_docx.parse_table_row("| a | b | c |") == ["a", "b", "c"]
        assert to_docx.parse_table_row("|x|y|") == ["x", "y"]
        assert to_docx.parse_table_row("| leading | trailing spaces |") == ["leading", "trailing spaces"]

    @pytest.mark.parametrize("line,expected", [
        # Good subtitle candidates
        ("Elasticity Engagement", True),
        ("Demand Forecasting Engagement", True),
        ("Q1 2026 Roadmap", True),
        # Rejected: ends with sentence punctuation
        ("This is a body sentence.", False),
        ("Final question?", False),
        ("Ending with exclaim!", False),
        ("Trailing colon:", False),
        # Rejected: too long
        ("Draft scoping document | Databricks FDE | Version v0.6 | 22 Apr 2026", False),
        # Rejected: too many words
        ("One two three four five six seven eight nine ten eleven words", False),
        # Rejected: lowercase start
        ("lowercase first letter", False),
        # Rejected: block markers
        ("## Not a subtitle", False),
        ("- bullet line", False),
        ("1. numbered item that looks like subtitle", False),
        ("| cell | table |", False),
        ("> quoted line", False),
        ("```", False),
        ("--------", False),
        ("ALL CAPS LINE", False),
        ("Status: Active", False),
        # Edge
        ("", False),
    ])
    def test_is_subtitle_candidate(self, line, expected):
        assert to_docx.is_subtitle_candidate(line) is expected

    @pytest.mark.parametrize("line,expected", [
        ("![alt text](image.png)", True),
        ("![](image.png)", True),
        ("![diagram](path/to/img.jpg)", True),
        ("![alt](img.png \"width=4in\")", True),
        ("  ![alt](img.png)  ", True),
        ("text ![alt](img.png) more", False),  # not a full-line image
        ("[link](url)", False),                 # missing ! prefix
        ("![alt](http://example.com/x.png)", True),
        ("not an image", False),
        ("", False),
    ])
    def test_is_image_block(self, line, expected):
        assert to_docx.is_image_block(line) is expected

    def test_parse_image_block_basic(self):
        assert to_docx.parse_image_block("![alt text](path.png)") == ("alt text", "path.png", None)

    def test_parse_image_block_with_title(self):
        assert to_docx.parse_image_block('![alt](path.png "width=4in")') == ("alt", "path.png", "width=4in")

    def test_parse_image_block_empty_alt(self):
        assert to_docx.parse_image_block("![](path.png)") == ("", "path.png", None)

    @pytest.mark.parametrize("line,expected", [
        ("```mermaid", True),
        ("``` mermaid", True),
        ("```Mermaid", True),       # case-insensitive
        ("```mermaid flowchart LR", True),  # extra args after language
        ("```", False),              # no language
        ("```python", False),        # different language
        ("```mermaidx", False),      # not an exact mermaid tag
        ("  ```mermaid", True),      # indented
        ("some text ```mermaid", False),  # not at start
    ])
    def test_is_mermaid_fence(self, line, expected):
        assert to_docx.is_mermaid_fence(line) is expected


# ============================================================
# HTML stripping
# ============================================================

class TestStripHtml:
    @pytest.mark.parametrize("text,expected", [
        ("plain text", "plain text"),
        ("<b>bold</b>", "bold"),
        ("text <i>italic</i> more", "text italic more"),
        ("<div class='x'>content</div>", "content"),
        ("<br/>", ""),
        ("<!-- comment -->", ""),
    ])
    def test_strip_html(self, text, expected):
        assert to_docx.strip_html(text) == expected


# ============================================================
# Integration tests — generate a .docx and verify structure
# ============================================================

def _convert_str(source: str, tmp_path: Path, title: str = "Test") -> Document:
    src = tmp_path / "input.txt"
    src.write_text(source)
    dst = tmp_path / "output.docx"
    to_docx.convert(src, dst, title)
    return Document(dst)


class TestIntegration:
    def test_title_appears_as_heading_0(self, tmp_path):
        doc = _convert_str("hello", tmp_path, title="My Document")
        assert doc.paragraphs[0].text == "My Document"
        assert doc.paragraphs[0].style.name == "Title"

    def test_title_line_not_duplicated_in_body(self, tmp_path):
        """When the first line of the file is also the title, it must not render twice."""
        src = "My Document\n\nBody content here.\n"
        doc = _convert_str(src, tmp_path, title="My Document")
        # Only the Title paragraph should have this text — no duplicate body paragraph
        matches = [p for p in doc.paragraphs if p.text == "My Document"]
        assert len(matches) == 1
        assert matches[0].style.name == "Title"

    def test_subtitle_from_header_block(self, tmp_path):
        """Second non-empty plain-text line after title renders as Subtitle style."""
        src = (
            "Dick's Sporting Goods (DSG)\n"
            "Elasticity Engagement\n"
            "Draft scoping document | Databricks FDE | v0.6 | 22 Apr 2026\n"
            "\n"
            "Body paragraph begins here.\n"
        )
        doc = _convert_str(src, tmp_path, title="Dick's Sporting Goods (DSG)")
        subtitles = [p for p in doc.paragraphs if p.style.name == "Subtitle"]
        assert len(subtitles) == 1
        assert subtitles[0].text == "Elasticity Engagement"
        # The meta line should still appear as body (not duplicated, not eaten)
        meta = [p for p in doc.paragraphs if "Draft scoping" in p.text]
        assert len(meta) == 1
        # And the body paragraph follows
        assert any("Body paragraph begins here" in p.text for p in doc.paragraphs)

    def test_no_subtitle_when_second_line_is_body_sentence(self, tmp_path):
        """A normal body sentence on line 2 must NOT be consumed as subtitle."""
        src = "Simple Title\n\nThis is a full body sentence.\n"
        doc = _convert_str(src, tmp_path, title="Simple Title")
        subtitles = [p for p in doc.paragraphs if p.style.name == "Subtitle"]
        assert len(subtitles) == 0
        assert any("full body sentence" in p.text for p in doc.paragraphs)

    def _make_tiny_png(self, path: Path):
        """Write a minimal valid 1x1 PNG for image-embedding tests."""
        # 67-byte minimal PNG (1x1, transparent)
        png_bytes = (
            b"\x89PNG\r\n\x1a\n"
            b"\x00\x00\x00\rIHDR"
            b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00"
            b"\x1f\x15\xc4\x89"
            b"\x00\x00\x00\rIDATx\x9cc\xfc\xcf\xc0P\x0f\x00\x04\x85\x01\x80"
            b"\x84\xa9\xf9c\x00\x00\x00\x00IEND\xaeB`\x82"
        )
        path.write_bytes(png_bytes)

    def test_image_block_embeds_picture(self, tmp_path):
        img = tmp_path / "test.png"
        self._make_tiny_png(img)
        src = f"![test image]({img})\n"
        doc = _convert_str(src, tmp_path)
        # python-docx exposes inline shapes (including images) via .inline_shapes
        assert len(doc.inline_shapes) == 1
        # Caption paragraph should carry the alt text in italic
        caption = [p for p in doc.paragraphs if p.text == "test image"]
        assert len(caption) == 1
        assert all(r.italic for r in caption[0].runs)

    def test_image_missing_file_renders_placeholder(self, tmp_path):
        src = "![oops](missing.png)\n"
        doc = _convert_str(src, tmp_path)
        assert len(doc.inline_shapes) == 0
        placeholder = [p for p in doc.paragraphs if "[image not found" in p.text]
        assert len(placeholder) == 1

    def test_mermaid_fallback_without_mmdc(self, tmp_path, monkeypatch):
        """When mmdc is unavailable, the Mermaid source is preserved as a code block
        with an italic note, rather than silently dropped."""
        monkeypatch.setattr(to_docx.shutil, "which", lambda name: None)
        src = (
            "```mermaid\n"
            "flowchart LR\n"
            "  A --> B\n"
            "```\n"
        )
        doc = _convert_str(src, tmp_path)
        # No image should be embedded
        assert len(doc.inline_shapes) == 0
        # Note paragraph should appear
        notes = [p for p in doc.paragraphs if "Mermaid diagram" in p.text]
        assert len(notes) == 1
        assert any(r.italic for r in notes[0].runs)
        # Mermaid source should appear as a code block
        code_paragraphs = [
            p for p in doc.paragraphs
            if any(r.font.name == "Consolas" for r in p.runs if r.font.name)
        ]
        code_text = "\n".join(p.text for p in code_paragraphs)
        assert "flowchart LR" in code_text
        assert "A --> B" in code_text

    def test_mermaid_renders_when_mmdc_present(self, tmp_path, monkeypatch):
        """Simulate mmdc being present and a successful render, verify image is embedded."""
        # Fake mmdc binary path
        monkeypatch.setattr(to_docx.shutil, "which",
                            lambda name: "/fake/mmdc" if name == "mmdc" else None)

        # Fake subprocess.run that writes a tiny valid PNG to the requested output path
        def fake_run(cmd, check=True, capture_output=True, timeout=60):
            # cmd is [mmdc, -i, in_path, -o, out_path, -b, transparent]
            out_path = Path(cmd[4])
            png_bytes = (
                b"\x89PNG\r\n\x1a\n"
                b"\x00\x00\x00\rIHDR"
                b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00"
                b"\x1f\x15\xc4\x89"
                b"\x00\x00\x00\rIDATx\x9cc\xfc\xcf\xc0P\x0f\x00\x04\x85\x01\x80"
                b"\x84\xa9\xf9c\x00\x00\x00\x00IEND\xaeB`\x82"
            )
            out_path.write_bytes(png_bytes)
            class Result: pass
            return Result()

        monkeypatch.setattr(to_docx.subprocess, "run", fake_run)

        src = (
            "```mermaid\n"
            "flowchart LR\n"
            "  A --> B\n"
            "```\n"
        )
        doc = _convert_str(src, tmp_path)
        # Image should be embedded
        assert len(doc.inline_shapes) == 1
        # No fallback note should appear
        notes = [p for p in doc.paragraphs if "Mermaid diagram" in p.text]
        assert len(notes) == 0

    def test_no_subtitle_when_second_line_is_heading(self, tmp_path):
        """An H2 on line 2 must NOT be consumed as subtitle."""
        src = "My Doc\n\n## Section One\n\nBody.\n"
        doc = _convert_str(src, tmp_path, title="My Doc")
        subtitles = [p for p in doc.paragraphs if p.style.name == "Subtitle"]
        assert len(subtitles) == 0
        h2s = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
        assert len(h2s) == 1 and h2s[0].text == "Section One"

    def test_explicit_heading_levels(self, tmp_path):
        src = "# H1\n\n## H2\n\n### H3\n\n#### H4\n"
        doc = _convert_str(src, tmp_path)
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert [(h.text, h.style.name) for h in headings] == [
            ("H1", "Heading 1"),
            ("H2", "Heading 2"),
            ("H3", "Heading 3"),
            ("H4", "Heading 4"),
        ]

    def test_all_caps_becomes_h1(self, tmp_path):
        doc = _convert_str("HELLO WORLD\n\nbody text\n", tmp_path)
        h1s = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert len(h1s) == 1
        assert h1s[0].text == "Hello World"

    def test_inline_bold(self, tmp_path):
        doc = _convert_str("This is **bold** text.", tmp_path)
        p = doc.paragraphs[1]  # [0] is title
        bold_runs = [r for r in p.runs if r.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "bold"

    def test_inline_italic(self, tmp_path):
        doc = _convert_str("This is *italic* text.", tmp_path)
        p = doc.paragraphs[1]
        italic_runs = [r for r in p.runs if r.italic]
        assert len(italic_runs) == 1
        assert italic_runs[0].text == "italic"

    def test_inline_bold_underscore(self, tmp_path):
        doc = _convert_str("This is __bold__ text.", tmp_path)
        p = doc.paragraphs[1]
        bold_runs = [r for r in p.runs if r.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "bold"

    def test_inline_italic_underscore(self, tmp_path):
        doc = _convert_str("This is _italic_ text.", tmp_path)
        p = doc.paragraphs[1]
        italic_runs = [r for r in p.runs if r.italic]
        assert len(italic_runs) == 1
        assert italic_runs[0].text == "italic"

    def test_inline_code(self, tmp_path):
        doc = _convert_str("Call `my_function` now.", tmp_path)
        p = doc.paragraphs[1]
        code_runs = [r for r in p.runs if r.font.name == "Consolas"]
        assert len(code_runs) == 1
        assert code_runs[0].text == "my_function"

    def test_inline_strikethrough(self, tmp_path):
        doc = _convert_str("This is ~~deleted~~ text.", tmp_path)
        p = doc.paragraphs[1]
        strike_runs = [r for r in p.runs if r.font.strike]
        assert len(strike_runs) == 1
        assert strike_runs[0].text == "deleted"

    def test_inline_link(self, tmp_path):
        doc = _convert_str("See [docs](https://docs.example.com) for info.", tmp_path)
        # Hyperlink is an OxmlElement appended to paragraph, not in .runs
        from docx.oxml.ns import qn
        p = doc.paragraphs[1]
        hyperlinks = p._p.findall(qn("w:hyperlink"))
        assert len(hyperlinks) == 1

    def test_html_stripped(self, tmp_path):
        doc = _convert_str("Text with <b>html</b> tags.", tmp_path)
        p = doc.paragraphs[1]
        full = "".join(r.text for r in p.runs)
        assert "<b>" not in full and "</b>" not in full
        assert "html" in full

    def test_bullets(self, tmp_path):
        doc = _convert_str("- first\n- second\n- **third** bold\n", tmp_path)
        bullets = [p for p in doc.paragraphs if p.style.name.startswith("List Bullet")]
        assert len(bullets) == 3
        assert bullets[0].text == "first"
        bold_in_third = [r for r in bullets[2].runs if r.bold]
        assert len(bold_in_third) == 1
        assert bold_in_third[0].text == "third"

    def test_numbered_list_ends_with_punctuation(self, tmp_path):
        doc = _convert_str(
            "1. First question?\n2. Second statement.\n3. Third colon:\n",
            tmp_path,
        )
        # All three are list items (end with punctuation, don't become headings)
        list_items = [p for p in doc.paragraphs if p.style.name == "List Number"]
        assert len(list_items) == 3

    def test_numbered_heading_no_punctuation(self, tmp_path):
        doc = _convert_str("1. Overview\n\nbody text\n", tmp_path)
        # "1. Overview" is short + no end punctuation -> H2
        h2s = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
        assert len(h2s) == 1
        assert h2s[0].text == "1. Overview"

    def test_dividers_dropped(self, tmp_path):
        doc = _convert_str("before\n\n========\n\nafter\n", tmp_path)
        texts = [p.text for p in doc.paragraphs if p.text]
        assert "========" not in texts

    def test_code_fence_block(self, tmp_path):
        src = "Before code\n\n```\ncode line 1\ncode line 2\n```\n\nAfter code\n"
        doc = _convert_str(src, tmp_path)
        code_paragraphs = [
            p for p in doc.paragraphs
            if any(r.font.name == "Consolas" for r in p.runs if r.font.name)
        ]
        assert len(code_paragraphs) == 2

    def test_pipe_table(self, tmp_path):
        src = (
            "| Role | Days |\n"
            "|---|---|\n"
            "| Lead | 36 |\n"
            "| Supporting | 20 |\n"
        )
        doc = _convert_str(src, tmp_path)
        assert len(doc.tables) == 1
        t = doc.tables[0]
        assert len(t.rows) == 3
        assert len(t.columns) == 2
        assert t.rows[0].cells[0].text == "Role"
        assert t.rows[2].cells[1].text == "20"
        # header row should have bold runs
        header_runs = t.rows[0].cells[0].paragraphs[0].runs
        assert any(r.bold for r in header_runs)

    def test_pipe_table_with_inline_bold(self, tmp_path):
        src = (
            "| Col | Val |\n"
            "|---|---|\n"
            "| a | **bold** |\n"
        )
        doc = _convert_str(src, tmp_path)
        cell = doc.tables[0].rows[1].cells[1]
        bold_runs = [r for r in cell.paragraphs[0].runs if r.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "bold"

    def test_blockquote(self, tmp_path):
        src = "regular\n\n> quoted line\n> continues\n\nafter\n"
        doc = _convert_str(src, tmp_path)
        quote_ps = [p for p in doc.paragraphs if "Quote" in p.style.name]
        assert len(quote_ps) == 1
        assert "quoted line" in quote_ps[0].text
        assert "continues" in quote_ps[0].text

    def test_kv_label(self, tmp_path):
        doc = _convert_str("Duration: 10 weeks\n", tmp_path)
        p = doc.paragraphs[1]
        bold_runs = [r for r in p.runs if r.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "Duration:"

    def test_mixed_inline_in_bullet(self, tmp_path):
        src = "- **Feature A** — uses `some_function` and [docs](https://ex.com)\n"
        doc = _convert_str(src, tmp_path)
        bullet = [p for p in doc.paragraphs if p.style.name.startswith("List Bullet")][0]
        bold = [r for r in bullet.runs if r.bold]
        code = [r for r in bullet.runs if r.font.name == "Consolas"]
        assert len(bold) == 1 and bold[0].text == "Feature A"
        assert len(code) == 1 and code[0].text == "some_function"

    def test_h1_dividers_together(self, tmp_path):
        """Replicates the actual drafts' section-title pattern."""
        src = (
            "========================================================\n"
            "1. OVERVIEW\n"
            "========================================================\n"
            "\n"
            "Body paragraph under overview.\n"
            "\n"
            "## Sub-section\n"
            "\n"
            "Another paragraph.\n"
        )
        doc = _convert_str(src, tmp_path)
        h1s = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
        h2s = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
        assert len(h1s) == 1
        assert h1s[0].text == "1. Overview"
        assert len(h2s) == 1
        assert h2s[0].text == "Sub-section"


# ============================================================
# Regression: things that previously broke
# ============================================================

class TestRegressions:
    def test_no_literal_asterisks_in_bold(self, tmp_path):
        """Inline bold markers must not appear as literal `**` in output."""
        doc = _convert_str("- **Cold-start handling** — description\n", tmp_path)
        bullet = doc.paragraphs[1]
        full_text = "".join(r.text for r in bullet.runs)
        assert "**" not in full_text

    def test_numbered_question_items_not_headings(self, tmp_path):
        """Enumerated questions (end with ?) must render as list items, not H2."""
        src = (
            "1. Which category?\n"
            "2. Who is the lead?\n"
            "3. What is the target accuracy level across multiple hierarchies and horizons?\n"
        )
        doc = _convert_str(src, tmp_path)
        list_items = [p for p in doc.paragraphs if p.style.name == "List Number"]
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        # 3 list items, zero sub-headings (the title uses "Title" style, not "Heading N")
        assert len(list_items) == 3
        assert len(headings) == 0
